# -*- coding: utf-8 -*-
"""生成股东版商业计划书与招商加盟计划书（Word .docx）"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.dirname(os.path.abspath(__file__))
GREEN = RGBColor(0x0E, 0x8A, 0x5A)
DARK = RGBColor(0x1F, 0x2D, 0x3D)


def set_font(run, name='微软雅黑', size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 0:
        set_font(p.add_run(text), size=22, bold=True, color=GREEN)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        set_font(p.add_run(text), size=16, bold=True, color=GREEN)
    elif level == 2:
        set_font(p.add_run(text), size=13, bold=True, color=DARK)
    else:
        set_font(p.add_run(text), size=11, bold=True, color=DARK)
    p.paragraph_format.space_before = Pt(10 if level else 24)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, size=11, bold=False, indent=False):
    p = doc.add_paragraph()
    set_font(p.add_run(text), size=size, bold=bold)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        set_font(hdr[i].paragraphs[0].add_run(h), size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '0E8A5A')
        tcPr.append(shd)
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = ''
            set_font(cells[i].paragraphs[0].add_run(str(v)), size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def build_shareholder_doc():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # 封面
    for _ in range(5):
        doc.add_paragraph()
    add_heading(doc, '健康管理加盟平台', 0)
    add_heading(doc, '商业计划书', 0)
    p = doc.add_paragraph()
    set_font(p.add_run('（股东版 · 内部资料）'), size=14, bold=True, color=DARK)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_font(p.add_run('以智能中医体质检测为入口 · AI 专业服务为引擎 · 加盟连锁为增长'), size=12, color=DARK)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_heading(doc, '一、项目概述', 1)
    add_heading(doc, '1.1 项目定位', 2)
    add_para(doc, '本项目是一套「总部 + 加盟门店」两级大健康 SaaS 平台。总部把「检测设备 + 管理系统 + AI 专业能力」打包赋能给养生馆、中医馆、美容院等线下门店；门店负责获客与本地服务，总部负责平台、AI、供应链与标准化，双方按服务分润。', indent=True)
    add_heading(doc, '1.2 核心价值', 2)
    add_para(doc, '1. 对门店：60 秒无创检测生成 40 份报告 + AI 解读，体验式获客、专业能力补强、数字化管理。', indent=True)
    add_para(doc, '2. 对总部：AI 服务为复购型、低边际成本收入，门店经营越活跃，总部收入越高。', indent=True)
    add_para(doc, '3. 对客户：无创快速了解自身体质，获得个性化调理方案与持续健康管理。', indent=True)

    add_heading(doc, '二、系统与技术资产', 1)
    add_table(doc,
              ['项目', '说明'],
              [
                  ['后端', 'NestJS + Prisma + SQLite，PM2 守护，阿里云 ECS 24 小时在线'],
                  ['桌面端', 'Electron + Vue3（Windows），安装即用，后续可扩展手机端/小程序'],
                  ['检测设备', 'Quantum Analyzer 手掌检测仪（USB HID，60 秒无创检测）'],
                  ['AI 能力', '已接入 7 家主流大模型（MiniMax/千问/豆包/Kimi/DeepSeek/智谱/蚂蚁百灵），总部自助配置'],
                  ['数据安全', '每日自动备份数据库（保留 7 份），门店数据严格隔离'],
              ],
              widths=[4, 12])

    add_heading(doc, '三、商业模式与盈利来源', 1)
    add_table(doc,
              ['收入项', '模式', '定价参考'],
              [
                  ['设备销售', '一次性销售检测仪', '¥3,980–¥9,800/台'],
                  ['软件订阅', '门店 SaaS 年费', '¥2,999–¥19,999/年'],
                  ['AI 服务分润', '门店付费申请 AI 服务', '报告解读 ¥9.9/次、调理方案 ¥99/次、远程会诊 ¥99/次'],
                  ['供应链集采', '理调耗材、养生品', '艾灸条、足浴包、药膳包等'],
                  ['加盟费/保证金', '品牌授权', '按城市分级 ¥1万–¥10万'],
              ],
              widths=[3.5, 5, 7.5])
    add_para(doc, '核心优势：AI 服务边际成本趋近于零，是可持续的复购型收入。', bold=True)

    add_heading(doc, '四、加盟体系设计', 1)
    add_table(doc,
              ['角色', '权限'],
              [
                  ['总部 admin', '加盟审批、门店管理、服务工单、订阅计费、AI 配置、全部门店报表'],
                  ['门店店长', '检测、客户、报告、订单、预约、调理方案、服务申请（仅本店数据）'],
                  ['门店店员', '按店长勾选的功能权限，细粒度控制'],
              ],
              widths=[3.5, 12.5])
    add_para(doc, '加盟流程：商家在线提交申请 → 总部审批 → 一键开通门店 + 商家账号 → 门店安装即用。')

    add_heading(doc, '五、财务测算（参考模型）', 1)
    add_heading(doc, '5.1 单店收入测算（月到店 200 人次，转化率 30%）', 2)
    add_table(doc,
              ['收入项', '测算', '月收入'],
              [
                  ['检测体验（60 单 × ¥99）', '60 × 99', '¥5,940'],
                  ['调理套餐（15 单 × ¥1999）', '15 × 1999', '¥29,985'],
                  ['年度管家（5 单 × ¥3999）', '5 × 3999', '¥19,995'],
                  ['合计', '', '¥55,920/月'],
              ],
              widths=[6, 5, 4])
    add_heading(doc, '5.2 50 家加盟门店年收入测算', 2)
    add_table(doc,
              ['收入项', '单店', '50 店合计（年）'],
              [
                  ['设备销售（一次性）', '¥6,000', '¥300,000'],
                  ['软件订阅年费', '¥9,999', '¥499,950'],
                  ['AI 服务分润（月均 100 次）', '¥990/月', '¥594,000'],
                  ['加盟费（一次性）', '¥30,000', '¥1,500,000'],
                  ['供应链（月均）', '¥500/月', '¥300,000'],
              ],
              widths=[6, 5, 5])
    add_para(doc, '注：以上为参考值，最终以实际定价与招商政策为准。', size=9)

    add_heading(doc, '六、近期计划与里程碑', 1)
    add_table(doc,
              ['阶段', '动作', '目标'],
              [
                  ['第 0–1 个月', '1 家自营店试点，跑通闭环，打磨话术与套餐', '验证单店模型'],
                  ['第 2–3 个月', '开放 3–5 家加盟，验证服务工单与分润', '跑通加盟模型'],
                  ['第 4–6 个月', '区域招商，标准化培训与供应链', '形成区域网络'],
                  ['第 7–12 个月', '多区域复制，数据看板驱动经营', '规模盈利'],
              ],
              widths=[3.5, 8, 4.5])

    add_heading(doc, '七、需要股东支持的决策', 1)
    add_para(doc, '1. 确认设备售价、加盟费、软件订阅年费的最终定价。')
    add_para(doc, '2. 确认首批试点城市与自营/加盟比例。')
    add_para(doc, '3. 确认是否需要融资用于设备备货与招商推广。')

    path = os.path.join(OUT, '健康管理加盟平台-商业计划书(股东版).docx')
    doc.save(path)
    return path


def build_franchise_doc():
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    for _ in range(5):
        doc.add_paragraph()
    add_heading(doc, '健康管理连锁平台', 0)
    add_heading(doc, '招商加盟计划书', 0)
    p = doc.add_paragraph()
    set_font(p.add_run('智能检测 + AI 专业服务 + 数字化管理'), size=14, bold=True, color=DARK)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_heading(doc, '一、市场机会', 1)
    add_para(doc, '大健康产业持续高速增长，养生馆、中医馆、美容院普遍面临获客难、留客难、专业人才稀缺、管理粗放三大痛点。我们以「智能检测 + AI 专业服务 + 数字化管理」三位一体，为门店提供可快速复制的整体解决方案。', indent=True)

    add_heading(doc, '二、平台对标与差异化', 1)
    add_table(doc,
              ['维度', '传统养生馆', '美团/点评', '本平台'],
              [
                  ['获客工具', '发传单/老客', '团购引流', '60秒检测+40份报告+AI解读，体验式获客'],
                  ['专业能力', '依赖老师傅', '无', 'AI 调理方案 + 远程医师会诊'],
                  ['客户管理', '手工记账', '无', '客户/预约/订单/报表一体化'],
                  ['复购机制', '弱', '弱', '检测套餐 + 疗程方案 + AI 报告持续触达'],
                  ['门店赋能', '无', '平台抽佣高', '总部系统赋能，按服务分润'],
              ],
              widths=[3, 4, 4, 6])
    add_para(doc, '对标结论：我们不是做流量平台，而是做门店经营系统 + 专业服务供应链，门店粘性更高、总部利润更厚。', bold=True)

    add_heading(doc, '三、加盟模式与费用', 1)
    add_table(doc,
              ['项目', '金额（参考）', '说明'],
              [
                  ['加盟费', '¥3万–¥10万', '按城市/区域分级'],
                  ['设备费', '¥3,980–¥9,800', 'Quantum Analyzer 检测仪'],
                  ['软件年费', '¥2,999–¥19,999', '按版本（基础/专业/旗舰）'],
                  ['保证金', '¥1万–¥5万', '可退'],
              ],
              widths=[4, 5, 7])
    add_para(doc, '分润机制：门店向总部申请 AI 服务按次付费（报告解读 ¥9.9、调理方案 ¥99、远程会诊 ¥99），总部统一配置 AI 与医师资源。')

    add_heading(doc, '四、加盟商收益测算（单店示例）', 1)
    add_table(doc,
              ['收入项', '测算', '月收入'],
              [
                  ['检测体验（60 单 × ¥99）', '60 × 99', '¥5,940'],
                  ['调理套餐（15 单 × ¥1999）', '15 × 1999', '¥29,985'],
                  ['年度管家（5 单 × ¥3999）', '5 × 3999', '¥19,995'],
                  ['合计', '', '¥55,920/月'],
              ],
              widths=[6, 5, 4])
    add_para(doc, '扣除房租人工与耗材后，单店毛利空间可观，加盟商 6–12 个月可回本（以实际经营为准）。', size=9)

    add_heading(doc, '五、总部八大支持', 1)
    for i, s in enumerate([
        '系统支持：总部统一部署，门店安装即用。',
        'AI 支持：调理方案、报告解读一键申请。',
        '远程会诊：复杂客户总部医师远程支持。',
        '培训支持：设备操作、检测话术、套餐销售培训。',
        '供应链支持：理调耗材集采，降低成本。',
        '营销支持：公众号模板消息、报告分享裂变。',
        '数据支持：经营报表，辅助门店决策。',
        '品牌支持：统一品牌形象与推广素材。',
    ], 1):
        add_para(doc, f'{i}. {s}')

    add_heading(doc, '六、招商政策与流程', 1)
    add_para(doc, '招商对象：现有养生馆、中医馆、美容院、健康管理中心，以及有健康行业经验的创业者、区域合伙人。')
    add_para(doc, '入驻流程：提交加盟申请 → 总部审核签约缴费 → 开通账号、设备发货 → 培训上线正式营业。')
    add_para(doc, '区域保护：按城市/商圈授予区域独家权，保障加盟商利益；首批优先开放 1–2 个试点城市。')

    add_heading(doc, '七、加盟保障与风险提示', 1)
    add_para(doc, '总部承诺系统 24 小时稳定运行，数据每日自动备份；门店数据独立隔离，客户隐私安全。加盟投资有风险，收益受门店选址、经营能力等因素影响，请理性评估。')

    path = os.path.join(OUT, '健康管理连锁平台-招商加盟计划书.docx')
    doc.save(path)
    return path


if __name__ == '__main__':
    p1 = build_shareholder_doc()
    p2 = build_franchise_doc()
    print('已生成:')
    print(' ', p1)
    print(' ', p2)
